import os
from PyPDF2 import PdfReader
import docx
from warnings import PendingDeprecationWarning, warn

def pdf_to_text(pdf_file):
    text = ""
    with open(pdf_file, 'rb') as f:
        reader = PdfReader(f)
        for page_num in range(len(reader.pages)):
            page_text = reader.pages[page_num].extract_text()
            if page_text:  # Check if text extraction was successful
                text += page_text + "\n\n"  # Add newlines to separate pages
            else:
                warn("Text extraction failed for page {}".format(page_num), PendingDeprecationWarning)
    return text

def pdf_to_docx(pdf_file, output_file):
    text = pdf_to_text(pdf_file)
    doc = docx.Document()
    for paragraph in text.split("\n\n"):
        doc.add_paragraph(paragraph)
    doc.save(output_file)

pdf_file = "./SQL_connectivity_in_Python1.pdf"
output_docx_file = "output_docx.docx"

pdf_to_docx(pdf_file, output_docx_file)






# import os
# from PyPDF2 import PdfReader
# import docx

# def pdf_to_text():
#     pdf_file = "./SQL connectivity in Python1.pdf"
#     text = ""
#     with open(pdf_file,'rb')as f:
#         reader = PdfReader(f)
#         for page_num in range(len(reader.pages)):
#             page_text = reader.pages[page_num].extract_text()
#             text += page_text
#         return text

# def pdf_to_docx(output_file):
#     text = pdf_to_text()
#     doc = docx.Document()
#     doc.add_paragraph(text)
#     doc.save(output_file)
    
# output_docx_file = "output_docx.docx"

# pdf_to_docx(output_docx_file)